"""Build the two public course DOCX files from their Markdown sources.

The public documents intentionally use a plain Chinese Word layout: A4 paper,
black text on white, restrained headings, simple tables/captions, and a centered
page number.  The two-page course manuscript follows the specified 9 pt
Microsoft YaHei body format; the long supporting draft uses 12 pt SimSun.
Both files are built directly from ``reports/sources`` so the Markdown remains
the single source of truth.

The final DOCX archives are normalized after saving.  Core properties and ZIP
entry timestamps are fixed so rebuilding an unchanged checkout is byte-stable.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
SOURCES = REPORTS / "sources"
PUBLIC = REPORTS / "public"
FIGURES = ROOT / "outputs" / "figures"

FORMAL_SOURCE = SOURCES / "formal_manuscript.md"
FULL_SOURCE = SOURCES / "full_analysis_draft.md"
FORMAL = PUBLIC / "今井达也_MLB调整决策_正式文稿.docx"
FULL = PUBLIC / "今井达也_MLB调整决策_完整分析底稿.docx"

FORMAL_BODY_FONT = "Microsoft YaHei"
FULL_BODY_FONT = "宋体"
HEADING_FONT = "黑体"
BLACK = "000000"
TABLE_HEADER_FILL = "D9D9D9"
FIXED_DATETIME = datetime(2026, 8, 12, 0, 0, 0)
FIXED_ZIP_TIME = (2026, 8, 12, 0, 0, 0)


def set_run_font(run, size: float, *, bold: bool = False, font: str = FULL_BODY_FONT) -> None:
    """Apply explicit black Chinese typography to a run."""
    run.font.name = font
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(BLACK)


def set_style_font(style, size: float, *, bold: bool = False, font: str = FULL_BODY_FONT) -> None:
    style.font.name = font
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(BLACK)


def configure_styles(
    doc: Document,
    *,
    body_size: float,
    body_font: str,
    line_spacing: float,
    compact: bool,
) -> None:
    """Resolve every used Word style to explicit plain-layout tokens."""
    normal = doc.styles["Normal"]
    set_style_font(normal, body_size, font=body_font)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2 if compact else 4)
    normal.paragraph_format.line_spacing = line_spacing

    title = doc.styles["Title"]
    set_style_font(title, 18 if compact else 20, bold=True, font=HEADING_FONT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4 if compact else 8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 12 if compact else 14, bold=True, font=HEADING_FONT)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(5 if compact else 8)
    subtitle.paragraph_format.line_spacing = 1.0

    heading_sizes = ({1: 12.5, 2: 11.5, 3: 10.5} if compact
                     else {1: 15, 2: 13, 3: 12})
    for level, size in heading_sizes.items():
        style = doc.styles[f"Heading {level}"]
        set_style_font(style, size, bold=True, font=HEADING_FONT)
        style.paragraph_format.space_before = Pt(8 if compact else 12)
        style.paragraph_format.space_after = Pt(3 if compact else 6)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    set_style_font(caption, 9, font=body_font)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(4 if compact else 6)
    caption.paragraph_format.line_spacing = 1.0


def add_page_number(section, *, size: float, font: str) -> None:
    """Add a centered PAGE field and leave the header empty."""
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.clear()

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)

    run = footer.add_run()
    set_run_font(run, size, font=font)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def setup_document(
    *,
    body_size: float,
    body_font: str,
    line_spacing: float,
    margins_cm: tuple[float, float, float, float],
    compact: bool,
) -> Document:
    doc = Document()
    configure_styles(
        doc,
        body_size=body_size,
        body_font=body_font,
        line_spacing=line_spacing,
        compact=compact,
    )
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(margins_cm[0])
        section.right_margin = Cm(margins_cm[1])
        section.bottom_margin = Cm(margins_cm[2])
        section.left_margin = Cm(margins_cm[3])
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.8)
        add_page_number(section, size=9, font=body_font)
    return doc


def add_inline_markup(
    paragraph,
    text: str,
    *,
    size: float,
    font: str,
    default_bold: bool = False,
) -> None:
    """Render the small Markdown subset used by the sources."""
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, bold=True, font=font)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, max(size - 0.5, 8.5), font=font)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, bold=default_bold, font=font)


def add_title(doc: Document, text: str, *, subtitle: bool = False) -> None:
    p = doc.add_paragraph(style="Subtitle" if subtitle else "Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    size = 12 if subtitle else 18
    if doc.styles["Normal"].font.size and doc.styles["Normal"].font.size.pt >= 12:
        size = 14 if subtitle else 20
    run = p.add_run(text)
    set_run_font(run, size, bold=True, font=HEADING_FONT)


def add_heading(doc: Document, text: str, *, level: int, compact: bool) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    size_map = ({1: 12.5, 2: 11.5, 3: 10.5} if compact
                else {1: 15, 2: 13, 3: 12})
    run = p.add_run(text)
    set_run_font(run, size_map[level], bold=True, font=HEADING_FONT)


def add_metadata(doc: Document, text: str, *, size: float, font: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text.rstrip())
    set_run_font(run, size, font=font)


def add_prose(
    doc: Document,
    text: str,
    *,
    size: float,
    font: str,
    line_spacing: float,
    compact: bool,
    indent: bool = True,
    bold: bool = False,
) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2 if compact else 4)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.first_line_indent = Pt(size * 2) if indent else Pt(0)
    p.paragraph_format.widow_control = True
    add_inline_markup(p, text, size=size, font=font, default_bold=bold)


def _next_numbering_id(numbering) -> tuple[int, int]:
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    return (max(abstract_ids, default=-1) + 1, max(num_ids, default=0) + 1)


def create_numbering(doc: Document, *, bullet: bool) -> int:
    """Create one real, restartable single-level Word list definition."""
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = _next_numbering_id(numbering)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend((tabs, ind))
    level.extend((start, fmt, lvl_text, suffix, p_pr))
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(
    doc: Document,
    text: str,
    *,
    num_id: int,
    size: float,
    font: str,
    line_spacing: float,
    compact: bool,
) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2 if compact else 4)
    p.paragraph_format.line_spacing = line_spacing
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    add_inline_markup(p, text, size=size, font=font)


def add_figure(
    doc: Document,
    path: Path,
    caption: str,
    *,
    width_cm: float,
    compact: bool,
    font: str,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c.add_run(caption)
    set_run_font(run, 9, font=font)


def _cm_to_dxa(value: float) -> int:
    return int(round(value / 2.54 * 1440))


def _set_cell_margins(cell, *, top: int = 60, start: int = 100, bottom: int = 60, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_cm: list[float]) -> None:
    """Apply fixed DXA geometry so Word and LibreOffice agree on widths."""
    widths_dxa = [_cm_to_dxa(width) for width in widths_cm]
    total_dxa = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(total_dxa))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "100")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "808080")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(widths_dxa[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _table_widths(column_count: int, usable_width_cm: float) -> list[float]:
    ratios = {
        2: [0.24, 0.76],
        3: [0.18, 0.50, 0.32],
        4: [0.14, 0.25, 0.31, 0.30],
    }.get(column_count, [1 / column_count] * column_count)
    widths = [round(usable_width_cm * ratio, 3) for ratio in ratios]
    widths[-1] = round(usable_width_cm - sum(widths[:-1]), 3)
    return widths


def add_markdown_table(
    doc: Document,
    rows: list[list[str]],
    *,
    size: float,
    font: str,
    usable_width_cm: float,
    line_spacing: float,
) -> None:
    if not rows or not rows[0]:
        return
    column_count = len(rows[0])
    if any(len(row) != column_count for row in rows):
        raise ValueError("Markdown table rows must have a consistent column count")
    table = doc.add_table(rows=len(rows), cols=column_count)
    _set_table_geometry(table, _table_widths(column_count, usable_width_cm))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), TABLE_HEADER_FILL)
                cell._tc.get_or_add_tcPr().append(shading)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = line_spacing
            add_inline_markup(
                paragraph,
                value,
                size=size,
                font=font,
                default_bold=row_index == 0,
            )
    spacer = doc.add_paragraph(style="Normal")
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    """Return table rows and the first unconsumed line, or ``None``."""
    if start + 1 >= len(lines):
        return None
    first = lines[start].strip()
    separator = lines[start + 1].strip()
    if not (first.startswith("|") and separator.startswith("|")):
        return None
    separator_cells = [cell.strip() for cell in separator.strip("|").split("|")]
    if not separator_cells or not all(re.fullmatch(r":?-{3,}:?", cell) for cell in separator_cells):
        return None
    rows = [[cell.strip() for cell in first.strip("|").split("|")]]
    index = start + 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def set_public_core_properties(doc: Document, *, title: str, subject: str) -> None:
    core = doc.core_properties
    core.title = title
    core.subject = subject
    core.author = "Anonymous"
    core.last_modified_by = "Anonymous"
    core.category = "Public research report"
    core.comments = "Public version; personal identifiers removed."
    core.keywords = "MLB, Statcast, pitching analysis"
    core.language = "zh-CN"
    core.identifier = "MLB-Analytics-public"
    core.revision = 1
    core.created = FIXED_DATETIME
    core.modified = FIXED_DATETIME


def normalize_docx_archive(path: Path) -> None:
    """Rewrite a DOCX with stable ordering, compression, and entry timestamps."""
    tmp = path.with_suffix(path.suffix + ".normalized")
    with ZipFile(path, "r") as source, ZipFile(tmp, "w", ZIP_DEFLATED, compresslevel=9) as target:
        for name in sorted(source.namelist()):
            info = ZipInfo(name, date_time=FIXED_ZIP_TIME)
            info.compress_type = ZIP_DEFLATED
            info.create_system = 3
            info.external_attr = 0o600 << 16
            target.writestr(info, source.read(name))
    tmp.replace(path)


def save_public_docx(doc: Document, path: Path, *, title: str, subject: str) -> None:
    set_public_core_properties(doc, title=title, subject=subject)
    path.parent.mkdir(parents=True, exist_ok=True)
    # Build beside the destination first.  Word may keep a submitted copy open
    # on Windows; an identical rebuild should still be verifiable without
    # attempting to overwrite that locked file.
    candidate = path.with_suffix(path.suffix + ".building")
    candidate.unlink(missing_ok=True)
    try:
        doc.save(candidate)
        normalize_docx_archive(candidate)
        if path.is_file() and path.read_bytes() == candidate.read_bytes():
            return
        candidate.replace(path)
    finally:
        candidate.unlink(missing_ok=True)


def build_formal() -> None:
    doc = setup_document(
        body_size=9,
        body_font=FORMAL_BODY_FONT,
        line_spacing=1.1,
        margins_cm=(1.9, 2.0, 1.9, 2.0),
        compact=True,
    )
    lines = FORMAL_SOURCE.read_text(encoding="utf-8").splitlines()
    current_list: int | None = None
    inserted_figure = False
    index = 0

    while index < len(lines):
        table_data = parse_markdown_table(lines, index)
        if table_data is not None:
            rows, index = table_data
            add_markdown_table(
                doc,
                rows,
                size=9,
                font=FORMAL_BODY_FONT,
                usable_width_cm=17.0,
                line_spacing=1.0,
            )
            current_list = None
            continue
        line = lines[index].strip()
        index += 1
        if not line:
            current_list = None
            continue
        if line.startswith("# "):
            add_title(doc, line[2:])
        elif line.startswith("## "):
            add_title(doc, line[3:], subtitle=True)
        elif line.startswith("### "):
            heading = line[4:]
            if heading.startswith("三、"):
                if not inserted_figure:
                    add_figure(
                        doc,
                        FIGURES / "figure_1_core_diagnosis.png",
                        "图1  NPB 2025与MLB 2026核心诊断。跨联盟球种价值仅比较方向。",
                        width_cm=15.5,
                        compact=True,
                        font=FORMAL_BODY_FONT,
                    )
                    inserted_figure = True
                doc.add_page_break()
            add_heading(doc, heading, level=1, compact=True)
            current_list = None
        elif line.startswith("公开版"):
            add_metadata(doc, line, size=9, font=FORMAL_BODY_FONT)
        elif re.match(r"^\d+\.\s+", line):
            if current_list is None:
                current_list = create_numbering(doc, bullet=False)
            text = re.sub(r"^\d+\.\s+", "", line)
            add_list_item(
                doc,
                text,
                num_id=current_list,
                size=9,
                font=FORMAL_BODY_FONT,
                line_spacing=1.1,
                compact=True,
            )
        elif line.startswith("数据来源："):
            add_prose(
                doc,
                line,
                size=9,
                font=FORMAL_BODY_FONT,
                line_spacing=1.0,
                compact=True,
                indent=False,
            )
            current_list = None
        else:
            add_prose(
                doc,
                line,
                size=9,
                font=FORMAL_BODY_FONT,
                line_spacing=1.1,
                compact=True,
                indent=not line.startswith("**"),
            )
            current_list = None

    save_public_docx(
        doc,
        FORMAL,
        title="今井达也应如何重返MLB先发轮值？",
        subject="MLB数据分析课程公开正式文稿",
    )


def build_full() -> None:
    doc = setup_document(
        body_size=12,
        body_font=FULL_BODY_FONT,
        line_spacing=1.5,
        margins_cm=(2.5, 2.5, 2.5, 2.5),
        compact=False,
    )
    lines = FULL_SOURCE.read_text(encoding="utf-8").splitlines()
    title_seen = False
    number_list: int | None = None
    bullet_list: int | None = None
    inserted_fig1 = False
    inserted_fig2 = False

    index = 0
    while index < len(lines):
        table_data = parse_markdown_table(lines, index)
        if table_data is not None:
            rows, index = table_data
            add_markdown_table(
                doc,
                rows,
                size=10.5,
                font=FULL_BODY_FONT,
                usable_width_cm=16.0,
                line_spacing=1.2,
            )
            number_list = None
            bullet_list = None
            continue
        line = lines[index].strip()
        index += 1
        if not line:
            number_list = None
            bullet_list = None
            continue
        if line.startswith("# ") and not title_seen:
            add_title(doc, line[2:])
            title_seen = True
        elif line.startswith("## 完整分析底稿"):
            add_title(doc, line[3:], subtitle=True)
        elif line.startswith("## "):
            add_heading(doc, line[3:], level=1, compact=False)
            number_list = None
            bullet_list = None
        elif line.startswith("### "):
            add_heading(doc, line[4:], level=2, compact=False)
            number_list = None
            bullet_list = None
        elif line.startswith(("公开版本：", "研究冻结日：", "项目目录：")):
            add_metadata(doc, line.rstrip(), size=10.5, font=FULL_BODY_FONT)
        elif line.startswith("- "):
            if bullet_list is None:
                bullet_list = create_numbering(doc, bullet=True)
            add_list_item(
                doc,
                line[2:],
                num_id=bullet_list,
                size=12,
                font=FULL_BODY_FONT,
                line_spacing=1.5,
                compact=False,
            )
            number_list = None
        elif re.match(r"^\d+\.\s+", line):
            if number_list is None:
                number_list = create_numbering(doc, bullet=False)
            text = re.sub(r"^\d+\.\s+", "", line)
            add_list_item(
                doc,
                text,
                num_id=number_list,
                size=12,
                font=FULL_BODY_FONT,
                line_spacing=1.5,
                compact=False,
            )
            bullet_list = None
        else:
            add_prose(
                doc,
                line,
                size=12,
                font=FULL_BODY_FONT,
                line_spacing=1.5,
                compact=False,
                indent=not line.startswith("**"),
            )
            number_list = None
            bullet_list = None

        if line.startswith("结论：保留 H1 与 H2") and not inserted_fig1:
            add_figure(
                doc,
                FIGURES / "figure_1_core_diagnosis.png",
                "图1  核心诊断：三振能力保留，控球与球种结构退化。",
                width_cm=15.5,
                compact=False,
                font=FULL_BODY_FONT,
            )
            inserted_fig1 = True
        if line.startswith("结论：牛棚是重新练习") and not inserted_fig2:
            add_figure(
                doc,
                FIGURES / "figure_2_role_platoon.png",
                "图2  角色与侧别：牛棚信号积极但样本有限，四缝线问题集中于左打。",
                width_cm=15.5,
                compact=False,
                font=FULL_BODY_FONT,
            )
            inserted_fig2 = True

    save_public_docx(
        doc,
        FULL,
        title="今井达也MLB调整决策完整分析底稿",
        subject="数据、计算、迭代、反证与决策门槛",
    )


def main() -> None:
    PUBLIC.mkdir(parents=True, exist_ok=True)
    build_formal()
    build_full()
    print("Built both public DOCX reports in reports/public/.")


if __name__ == "__main__":
    main()
